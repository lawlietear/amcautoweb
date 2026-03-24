# -*- coding: utf-8 -*-
"""
提取 Word 模板中的 Jinja2 变量
"""
import re
import sys
from docx import Document


def extract_jinja2_variables(text):
    """从文本中提取 {{ variable }} 格式的变量"""
    # 匹配 {{ variable }} 或 {{ variable.attribute }} 格式
    pattern = r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}'
    matches = re.findall(pattern, text)
    return matches


def extract_from_paragraph(para):
    """从段落中提取变量"""
    variables = []
    
    # 直接提取段落文本中的变量
    text = para.text
    variables.extend(extract_jinja2_variables(text))
    
    # 检查每个 run（处理可能的分段情况）
    full_text = ""
    for run in para.runs:
        full_text += run.text
    
    if full_text != text:
        variables.extend(extract_jinja2_variables(full_text))
    
    return list(set(variables))  # 去重


def extract_from_table(table):
    """从表格中提取变量"""
    variables = []
    for row in table.rows:
        for cell in row.cells:
            # 提取单元格中所有段落的变量
            for para in cell.paragraphs:
                variables.extend(extract_from_paragraph(para))
            # 递归处理嵌套表格
            for nested_table in cell.tables:
                variables.extend(extract_from_table(nested_table))
    return list(set(variables))


def extract_variables_from_docx(file_path):
    """从 docx 文件中提取所有变量"""
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"错误：无法打开文件 {file_path}")
        print(f"详情：{str(e)}")
        return []
    
    all_variables = []
    
    # 提取段落中的变量
    print(f"\n📄 扫描段落 ({len(doc.paragraphs)} 个)...")
    for i, para in enumerate(doc.paragraphs):
        vars_in_para = extract_from_paragraph(para)
        if vars_in_para:
            all_variables.extend(vars_in_para)
            print(f"  段落 {i+1}: {vars_in_para}")
    
    # 提取表格中的变量
    print(f"\n📊 扫描表格 ({len(doc.tables)} 个)...")
    for i, table in enumerate(doc.tables):
        vars_in_table = extract_from_table(table)
        if vars_in_table:
            all_variables.extend(vars_in_table)
            print(f"  表格 {i+1}: {vars_in_table}")
    
    # 去重并排序
    unique_variables = sorted(list(set(all_variables)))
    
    return unique_variables


def suggest_field_config(var_name):
    """根据变量名推测字段配置"""
    # 金额字段
    if any(keyword in var_name.lower() for keyword in ['amount', 'price', 'total', 'principal', 'interest', 'fee', 'cost']):
        return {
            "type": "numeric",
            "format": "thousand",
            "required": True
        }
    
    # 日期字段
    if any(keyword in var_name.lower() for keyword in ['date', 'time', 'day']):
        return {
            "type": "text",
            "required": True
        }
    
    # 数字字段
    if any(keyword in var_name.lower() for keyword in ['count', 'number', 'seq', 'no', 'rate']):
        return {
            "type": "text",
            "required": True
        }
    
    # 默认文本
    return {
        "type": "text",
        "required": True
    }


if __name__ == '__main__':
    # 设置 UTF-8 编码
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    
    if len(sys.argv) < 2:
        # 默认读取收购业务批复模板
        file_path = r'autodocweb_v2/word_templates/template_收购业务批复.docx'
    else:
        file_path = sys.argv[1]
    
    print(f"🔍 正在分析模板：{file_path}")
    print("=" * 60)
    
    variables = extract_variables_from_docx(file_path)
    
    print("\n" + "=" * 60)
    print(f"✅ 发现 {len(variables)} 个变量：")
    print("=" * 60)
    
    for var in variables:
        config = suggest_field_config(var)
        req_mark = "⭐" if config['required'] else ""
        print(f"  - {var} [{config['type']}]{req_mark}")
    
    # 生成配置文件片段
    print("\n" + "=" * 60)
    print("📋 建议的 JSON 配置（复制到 config_risk_compliance.json）：")
    print("=" * 60)
    
    import json
    config_list = []
    for var in variables:
        config = suggest_field_config(var)
        field_config = {
            "key": var,
            "label": var.replace('_', ' ').title(),
            "type": config["type"],
            "required": config["required"]
        }
        if config.get("format"):
            field_config["format"] = config["format"]
        
        # 批复日期特殊处理
        if var == 'approval_date':
            field_config["readonly"] = True
            field_config["auto"] = "current_date"
            field_config["required"] = False
            field_config["description"] = "自动生成：当前日期"
        
        config_list.append(field_config)
    
    print(json.dumps(config_list, ensure_ascii=False, indent=2))
